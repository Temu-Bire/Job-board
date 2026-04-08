from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
README = ROOT / "README.md"
OUT = ROOT / "CareerConnect_Documentation.docx"


def is_hr(line: str) -> bool:
    return line.strip() in {"---", "***", "___"}


def parse_md_table(lines: list[str], start_index: int):
    """Parse a GitHub-style markdown table starting at start_index.
    Returns: (table_rows, next_index) or (None, start_index) if not a table."""
    if start_index + 1 >= len(lines):
        return None, start_index
    header = lines[start_index].strip()
    sep = lines[start_index + 1].strip()
    if "|" not in header or "|" not in sep:
        return None, start_index
    if not re.match(r"^\|?[\s:\-|]+\|?$", sep):
        return None, start_index

    def split_row(row: str):
        # trim leading/trailing pipes and split
        row = row.strip().strip("|")
        return [c.strip() for c in row.split("|")]

    rows = [split_row(header)]
    i = start_index + 2
    while i < len(lines) and "|" in lines[i]:
        row = lines[i].strip()
        if not row:
            break
        rows.append(split_row(row))
        i += 1
    return rows, i


def add_code_block(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def main():
    if not README.exists():
        raise SystemExit(f"README not found: {README}")

    md = README.read_text(encoding="utf-8").splitlines()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []

    while i < len(md):
        line = md[i]

        # fenced code blocks
        if line.strip().startswith("```"):
            fence = line.strip()
            if not in_code:
                in_code = True
                code_lang = fence.replace("```", "").strip()
                code_lines = []
            else:
                in_code = False
                block = "\n".join(code_lines).rstrip()
                if code_lang.lower() == "mermaid":
                    doc.add_paragraph("Mermaid diagram (rendered in Markdown):")
                add_code_block(doc, block)
                code_lang = ""
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        # horizontal rule
        if is_hr(line):
            doc.add_paragraph("")
            i += 1
            continue

        # markdown table
        rows, next_i = parse_md_table(md, i)
        if rows:
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx in range(cols):
                    table.cell(r_idx, c_idx).text = row[c_idx] if c_idx < len(row) else ""
            i = next_i
            continue

        # bullets (simple)
        stripped = line.strip()
        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            i += 1
            continue

        # blank line
        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        # normal paragraph (keep simple; keep markdown syntax as text)
        doc.add_paragraph(line)
        i += 1

    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()

